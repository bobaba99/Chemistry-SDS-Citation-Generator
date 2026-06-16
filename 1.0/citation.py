#!/usr/bin/env python3
"""Generate SDS citations from safety data sheet PDFs.

Every ``*.pdf`` in the current directory has its first page parsed into a
structured record. The records are written out in two formats:

* ``export.docx`` -- a numbered, ACS-style reference list, and
* ``export.ris``  -- an RIS file that imports into Zotero, EndNote,
  Mendeley and other reference managers.

Both Sigma-Aldrich and Thermo Fisher / Alfa Aesar safety data sheets are
supported; the supplier is detected automatically from the document text.

ACS citation format::

    chemical name (italic); CAS RN; catalogue/product number, revision;
    supplier: city, state, revision date. URL (accessed date)
"""
import os
import re
import calendar
from datetime import datetime

from PyPDF2 import PdfReader
from docx import Document

DOCX_EXPORT = 'export.docx'
RIS_EXPORT = 'export.ris'


def first_page_text(pdf_path):
    """Return the extracted text of the first page of *pdf_path*."""
    reader = PdfReader(pdf_path)
    return reader.pages[0].extract_text() or ''


def search(pattern, text, field, flags=0):
    """Return the first capture group of *pattern*, or raise a clear error."""
    match = re.search(pattern, text, flags)
    if not match:
        raise ValueError('could not find %s' % field)
    return match.group(1).strip()


def accessed_today():
    return datetime.today().strftime('%Y-%m-%d')


def sigma_city(text):
    """Extract ``City, ST`` from the Sigma-Aldrich supplier address block."""
    block = text.split('Company', 1)[-1].split('Telephone', 1)[0]
    match = re.search(r"([A-Z][A-Z .'-]+?)\s+([A-Z]{2})\s{2,}\w", block)
    if not match:
        raise ValueError('could not find supplier city/state')
    return '%s, %s' % (match.group(1).strip().title(), match.group(2))


def parse_thermo_date(raw):
    for fmt in ('%d-%B-%Y', '%d-%b-%Y'):
        try:
            return datetime.strptime(raw, fmt)
        except ValueError:
            continue
    raise ValueError('unrecognised revision date: %s' % raw)


def sigmaaldrich(text):
    """Parse a Sigma-Aldrich SDS into a citation record."""
    date_match = re.search(r'Revision Date\s+(\d{2})\.(\d{2})\.(\d{4})', text)
    if not date_match:
        raise ValueError('could not find revision date')
    day, month, year = int(date_match.group(1)), int(date_match.group(2)), date_match.group(3)
    number = search(r'Product Number\s*:?\s*(\S+)', text, 'product number')

    return {
        'name': search(r'Product name\s*:\s*(.+)', text, 'product name'),
        'cas': search(r'CAS-No\.?\s*:?\s*(\d+-\d+-\d+)', text, 'CAS number'),
        'number': number,
        'version': search(r'Version\s+(\d+\.\d+)', text, 'version'),
        'supplier': 'Sigma-Aldrich',
        'publisher': 'Sigma-Aldrich',
        'place': sigma_city(text),
        'date_display': '%s %d, %s' % (calendar.month_name[month], day, year),
        'date_iso': '%s/%02d/%02d' % (year, month, day),
        'year': year,
        'url': ('https://www.sigmaaldrich.com/MSDS/MSDS/DisplayMSDSPage.do'
                '?country=CA&language=en&productNumber=' + number
                + '&brand=SIAL&PageToGoToURL=%2Fsafety-center.html'),
    }


def thermofisher(text):
    """Parse a Thermo Fisher / Alfa Aesar SDS into a citation record."""
    raw_date = search(r'Revision Date\s+(\d{1,2}-[A-Za-z]+-\d{4})', text, 'revision date')
    date = parse_thermo_date(raw_date)
    number = search(r'Cat No\.?\s*:?\s*(\S+)', text, 'catalogue number')

    return {
        'name': search(r'Product Name\s+(.+)', text, 'product name'),
        'cas': search(r'CAS-No\.?\s*:?\s*(\d+-\d+-\d+)', text, 'CAS number'),
        'number': number,
        'version': search(r'Revision Number\.?\s*:?\s*(\d+)', text, 'revision number'),
        'supplier': 'Alfa Aesar, Thermo Fisher',
        'publisher': 'Thermo Fisher Scientific',
        'place': 'Ward Hill, MA',
        'date_display': date.strftime('%B %d, %Y'),
        'date_iso': date.strftime('%Y/%m/%d'),
        'year': date.year,
        'url': 'https://www.alfa.com/en/msds/?language=CE&subformat=AGHS&sku=' + number,
    }


def parse_pdf(pdf_path):
    """Detect the supplier from the document text and return its record."""
    text = first_page_text(pdf_path)
    if 'Cat No' in text:
        return thermofisher(text)
    if 'Product Number' in text:
        return sigmaaldrich(text)
    raise ValueError('unrecognised SDS format '
                     '(not Sigma-Aldrich or Thermo Fisher)')


def add_citation(doc, record):
    """Append one numbered ACS citation to *doc* (chemical name in italics)."""
    para = doc.add_paragraph(style='List Number')
    para.add_run(record['name']).italic = True
    para.add_run(
        '; CAS RN: %(cas)s; %(number)s, rev. %(version)s; '
        '%(supplier)s: %(place)s, %(date_display)s. %(url)s' % record
        + ' (accessed %s)' % accessed_today()
    )


def ris_entry(record):
    """Render one record as an RIS reference (report type)."""
    fields = [
        ('TY', 'RPRT'),
        ('TI', '%s: Safety Data Sheet' % record['name']),
        ('AU', record['publisher']),
        ('PB', record['publisher']),
        ('CY', record['place']),
        ('PY', str(record['year'])),
        ('DA', record['date_iso']),
        ('ET', 'rev. %s' % record['version']),
        ('M1', record['number']),
        ('UR', record['url']),
        ('KW', 'CAS RN %s' % record['cas']),
        ('N1', 'CAS Registry Number: %s. Accessed %s.' % (record['cas'], accessed_today())),
        ('ER', ''),
    ]
    return '\r\n'.join('%s  - %s' % (tag, value) for tag, value in fields) + '\r\n'


def main():
    pdfs = sorted(f for f in os.listdir('.') if f.lower().endswith('.pdf'))
    if not pdfs:
        print('No PDF files found in the current directory.')
        return

    records = []
    for pdf in pdfs:
        try:
            records.append(parse_pdf(pdf))
            print('Added citation for %s' % pdf)
        except Exception as exc:
            print('Skipped %s: %s' % (pdf, exc))

    if not records:
        print('No citations were generated.')
        return

    doc = Document(DOCX_EXPORT) if os.path.isfile(DOCX_EXPORT) else Document()
    for record in records:
        add_citation(doc, record)
    doc.save(DOCX_EXPORT)

    with open(RIS_EXPORT, 'w', encoding='utf-8', newline='') as ris:
        ris.write('\r\n'.join(ris_entry(record) for record in records))

    print('Wrote %d citation(s) to %s and %s'
          % (len(records), DOCX_EXPORT, RIS_EXPORT))


if __name__ == '__main__':
    main()
